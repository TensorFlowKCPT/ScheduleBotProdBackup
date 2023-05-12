import docx
import os
from ScheduleClass import Schedule
from UrokClass import Urok
import datetime
import regex

def GetWordSchedule(Data:str,group:str):
    date_format = '%d.%m.%Y'
    Date = datetime.datetime.strptime(Data, date_format)

    path = "ExcelAndWord/Change.docx"
    doc = docx.Document(path)

    parag = doc.paragraphs

    lst = []
    for i in parag:
        if "на" in i.text:
            lst.append(i.text)
    
    arrdate = Data.split(".")[0]

    indx = None
    for i in range(len(lst)):
        if arrdate in lst[i]:
            indx = i
    try:
        table = doc.tables[indx]
    except:
        Urokq = []
        return[Schedule(Date, Urokq),True]


    arr = []
    for row in table.rows:
        row_values = [cell.text for cell in row.cells]
        arr.append(row_values)

    izmen = []
    for i in arr:
        if group in i[0]:
            for j in i:
                izmen.append(j.split("\n"))
    try:
        result = preobrazovatel(izmen)
    except:
            Urokq = []
            return[Schedule(Date, Urokq),True]
    urok_list = []

    urokspat = regex.compile(r"\d+-\d+")
    for i, urok_info in enumerate(result['Урок']):
        try:
            urok_number = urok_info[0] if len(urok_info) == 1 else '-'.join(urok_info)
            if regex.match(urokspat,urok_number):
                start, end = urok_number.split('-')
                for j in range(int(start), int(end) + 1):
                    
                    urok = Urok(Number=j, Name=result['Изменения'][i], Kabinet='', Prepod=result['Преподаватель'][i])
                    try: 
                        urok.Kabinet = result['Кабинет'][i]
                    except IndexError:
                        urok.Kabinet = urok_list[-1].Kabinet

                    urok_list.append(urok)
            else:     
                urok = Urok(Number=urok_number[0], Name=result['Изменения'][i], Kabinet='', Prepod=result['Преподаватель'][i])
                try: 
                    urok.Kabinet = result['Кабинет'][i]
                except IndexError:
                    urok.Kabinet = urok_list[-1].Kabinet
                urok_list.append(urok)
        except:
            return [Schedule(Date, urok_list),True]
    
    return [Schedule(Date, urok_list),False]



def preobrazovatel(lst:list):
    arg = {"Урок":[], "Изменения": [], "Преподаватель": [],"Кабинет": []}
    #Обработка уроков
    while ' ' in lst[1]:
           lst[1].remove(' ')
    while '' in lst[1]:
           lst[1].remove('')
    urokspat = regex.compile(r"\d+-\d+|\d+")
    for i in lst[1]:
        matches = urokspat.findall(i)
        if len(matches) > 1:
            start, end = map(int, matches[0].split('-'))
            arg['Урок'].append([str(j) for j in range(start, end + 1)])
        elif len(matches) == 1:
            arg['Урок'].append([matches[0]])

    #Первичная обработка кабинетов
    
    for i in lst[4]:
          if i==' ' or i=='':
            lst[4].remove(i)
    #Обработка Названия и преподов
    while ' ' in lst[3]:
           lst[3].remove(' ')
    while 'Будет' in lst[3]:
           lst[3].remove('Будет')
    while 'Замена' in lst[3]:
           lst[3].remove('Замена')
    
    pat = r"\b([А-ЯЁ][а-яё]+)\s+([А-ЯЁ]\.[А-ЯЁ]\.)"
    pattern = regex.compile(pat)

    for i, s in enumerate(lst[3]):
        match = pattern.match(s)
        if match:
            lst[3][i - 1] += "*" + lst[3][i]
            lst[3].pop(i)
    while "" in lst[3]:
        lst[3].remove("")
    pat = r"Отмена"
    re = regex.compile(pat)
    
    i = 0
    while i < len(lst[3]):
        if re.search(lst[3][i]):
            arg['Урок'].pop(i)
            lst[3].pop(i)
            lst[4].pop(i)
        else:
            i += 1
    #Вторичная обработка кабинетов и их внос
    pat2 = r"\d"
    for i in lst[4]:
        if regex.match(pat2,i.strip()):
            arg['Кабинет'].append(int(i.strip()))
        elif i.strip()==('с/з') or i.strip()==('ч/з'):
             arg['Кабинет'].append(i)
    
    #Конец(Надеюсь)
    for i in lst[3]:
        if 'Переносится' in i:
            arg['Урок'].pop(lst[3].index(i))
        if i == 'Замена кабинета':
            arg['Изменения'].append('Замена кабинета')
            arg['Преподаватель'].append('Замена кабинета')
        else:
            try:
                arg['Изменения'].append(i.split('*')[0])
                arg['Преподаватель'].append(i.split('*')[1])
            except:
                 if lst[3].index(i)%2==0:
                    arg['Изменения'].append(i)
                 else:
                     arg['Преподаватель'].append(i)
    
    #Перевыкид на случай импосторов
    Budet = r'Будет'
    Zamena = r'^Замена$'
    for key in ['Изменения', 'Преподаватель']:
        arg[key] = [line for line in arg[key] if not regex.search(f'{Budet}|{Zamena}', line)]
    temp_ = []
    for i in arg['Изменения']:
        if not 'Переносится' in i:
            temp_.append(i)
    arg['Изменения'] = temp_
    return arg

parsresult = GetWordSchedule("12.05.2023","ИСиП 22-11-1")
if(len(parsresult[0].Uroki)>0):
   for i in parsresult[0].Uroki:
      print(i.Number, i.Name, i.Prepod, i.Kabinet)
else:
   print("Изменений нет")
